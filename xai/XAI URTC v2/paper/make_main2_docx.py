"""Build main2.docx: the manuscript rebuilt natively in the IEEE conference
Word template's structure and typography (Times New Roman throughout, 24pt
title, 11pt author names, 9pt bold abstract, 10pt two-column body, 8pt
captions and references, small-caps section heads).

Unlike main.docx (a PDF-reflow conversion that relies on text frames and
breaks in Google Docs), this document is plain flowed content: one
single-column section for title and authors, then a true two-column
section, ordinary tables, and inline images. It imports cleanly into
Google Docs.

Reads main.tex, main.aux (citation and label numbers), main.bbl (formatted
references), and ../figures/*.png. Run from paper/: python make_main2_docx.py
Style values are taken from the official conference-template-letter.docx.
"""

import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEX = open("main.tex", encoding="utf-8").read()
AUX = open("main.aux", encoding="utf-8").read()
BBL = open("main.bbl", encoding="utf-8").read()

FONT = "Times New Roman"
COL_W = Pt(252)          # one column: (612 - 2*44.65 - 18) / 2 approx.
FULL_W = Pt(522.7)       # both columns
# The IEEE template mandates em dash lead-ins for Abstract and Index Terms.
DASH = "\u2014"
NBSP = "\u00a0"

CITE = {k: n for k, n in re.findall(r"\\bibcite\{([^}]+)\}\{(\d+)\}", AUX)}
LABEL = {k: n for k, n in re.findall(r"\\newlabel\{([^}]+)\}\{\{([^}]*)\}\{[^}]*\}", AUX)}

GREEK = {
    r"\rho": "\u03c1", r"\phi": "\u03c6", r"\Phi": "\u03a6",
    r"\Delta": "\u0394", r"\mu": "\u03bc", r"\sum": "\u03a3",
    r"\times": "\u00d7", r"\in": "\u2208",
}


def math_runs(m, base=None):
    """Render a $...$ fragment as run specs: letters italic, ^/_ scripts."""
    out = []
    base = base or {}
    for k, v in GREEK.items():
        m = m.replace(k, v)
    m = (m.replace(r"\max", "\x00max\x00").replace(r"\bigl", "").replace(r"\bigr", "")
           .replace(r"\left", "").replace(r"\right", "").replace(r"\,", " ").replace(r"\;", " "))
    i = 0
    while i < len(m):
        c = m[i]
        if c in "^_":
            key = "sup" if c == "^" else "sub"
            if i + 1 < len(m) and m[i + 1] == "{":
                j = m.index("}", i + 2)
                frag, i = m[i + 2:j], j + 1
            else:
                frag, i = m[i + 1], i + 2
            for r in math_runs(frag, base):
                r[key] = True
                out.append(r)
        elif c == "\x00":
            j = m.index("\x00", i + 1)
            out.append(dict(base, t=m[i + 1:j]))
            i = j + 1
        elif c.isalpha():
            j = i
            while j < len(m) and m[j].isalpha():
                j += 1
            out.append(dict(base, t=m[i:j], i=True))
            i = j
        elif c in "{}":
            i += 1
        else:
            out.append(dict(base, t=c))
            i += 1
    return out


def tex_runs(s, base=None):
    """Convert inline LaTeX text to run specs."""
    base = base or {}
    s = re.sub(r"\\label\{[^}]*\}", "", s)
    s = re.sub(r"\\cite\{([^}]+)\}",
               lambda m: ", ".join("[%s]" % CITE[k.strip()] for k in m.group(1).split(",")), s)
    s = re.sub(r"\\eqref\{([^}]+)\}", lambda m: "(%s)" % LABEL[m.group(1)], s)
    s = re.sub(r"\\ref\{([^}]+)\}", lambda m: LABEL[m.group(1)], s)
    s = s.replace(r"\ugm{}", "$\\mu$g/m$^3$").replace(r"\ugm", "$\\mu$g/m$^3$")
    s = (s.replace(r"\%", "%").replace(r"\&", "&").replace("``", "\u201c")
          .replace("''", "\u201d").replace("~", NBSP))
    s = re.sub(r"\s+", " ", s)
    out = []
    pos = 0
    for m in re.finditer(r"\$([^$]+)\$|\\emph\{([^}]+)\}", s):
        if m.start() > pos:
            out.append(dict(base, t=s[pos:m.start()]))
        if m.group(1) is not None:
            out.extend(math_runs(m.group(1), base))
        else:
            out.extend(tex_runs(m.group(2), dict(base, i=True)))
        pos = m.end()
    if pos < len(s):
        out.append(dict(base, t=s[pos:]))
    return [r for r in out if r.get("t")]


def bbl_runs(txt):
    """Clean one .bbl entry: protect math and emph, strip grouping braces."""
    maths = []

    def hold(m):
        maths.append(m.group(0))
        return "\x04%d\x04" % (len(maths) - 1)

    txt = re.sub(r"\$[^$]*\$", hold, txt)
    txt = re.sub(r"\\emph\{([^}]*)\}", "\x02\\1\x03", txt)
    txt = txt.replace("{", "").replace("}", "")
    txt = re.sub(r"\x04(\d+)\x04", lambda m: maths[int(m.group(1))], txt)
    runs = []
    italic = False
    for part in re.split(r"([\x02\x03])", txt):
        if part == "\x02":
            italic = True
        elif part == "\x03":
            italic = False
        elif part:
            runs.extend(tex_runs(part, {"i": True} if italic else None))
    return runs


def add_runs(p, runs, size):
    for r in runs:
        run = p.add_run(r["t"])
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.italic = bool(r.get("i"))
        run.font.bold = bool(r.get("b"))
        run.font.superscript = bool(r.get("sup"))
        run.font.subscript = bool(r.get("sub"))
        run.font.small_caps = bool(r.get("sc"))
    return p


def para(doc, runs, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None,
         before=0, after=0, line=None, exact=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    if indent is not None:
        pf.first_line_indent = Pt(indent)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line) if exact else line
    add_runs(p, runs, size)
    return p


def set_cols(section, n, space_pt):
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(int(space_pt * 20)))


def page_setup(section):
    section.page_width = Pt(612)
    section.page_height = Pt(792)
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(44.65)
    section.right_margin = Pt(44.65)


def cell_border(tcPr, edge, sz):
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    b = OxmlElement("w:" + edge)
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:color"), "000000")
    borders.append(b)


doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)
sec = doc.sections[0]
page_setup(sec)

body = TEX.split(r"\begin{document}")[1]

# Title (template style: paper title, 24pt centered).
title = re.search(r"\\title\{(.*?)\}", body, re.S).group(1)
para(doc, tex_runs(re.sub(r"\s+", " ", title).strip()), size=24,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

# Author block: borderless 1x3 table (renders like the template's author
# columns but survives the Google Docs importer).
AUTHORS = [
    ["Saketh Chebrolu", "St. Mark's School of Texas", "Dallas, TX, USA",
     "chebrolusaketh@gmail.com"],
    ["Nathan Tan", "St. Mark's School of Texas", "Dallas, TX, USA",
     "nathantan2027@gmail.com"],
    ["Yifeng Wang", "School of Civil and Environmental Engineering",
     "Georgia Institute of Technology", "Atlanta, GA, USA",
     "ywang3627@gatech.edu"],
]
atable = doc.add_table(rows=1, cols=3)
atable.style = doc.styles["Normal Table"]
for cell, block in zip(atable.rows[0].cells, AUTHORS):
    first = True
    for k, line in enumerate(block):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(18 if k == 0 else 0)
        pf.space_after = Pt(2 if k == 0 else 0)
        add_runs(p, tex_runs(line), 11 if k == 0 else 10)

# Two-column body section (template: 2 cols, 18pt gap).
sec2 = doc.add_section(WD_SECTION.CONTINUOUS)
page_setup(sec2)
set_cols(sec2, 2, 18)

abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", body, re.S).group(1)
p = para(doc, [], size=9, indent=13.6, after=10)
add_runs(p, [{"t": "Abstract" + DASH, "i": True, "b": True}], 9)
add_runs(p, [dict(r, b=True) for r in tex_runs(re.sub(r"\s+", " ", abstract).strip())], 9)

keywords = re.search(r"\\begin\{IEEEkeywords\}(.*?)\\end\{IEEEkeywords\}", body, re.S).group(1)
p = para(doc, [], size=9, indent=13.7, after=6)
add_runs(p, [{"t": "Index Terms" + DASH, "i": True, "b": True}], 9)
add_runs(p, [dict(r, b=True, i=True) for r in tex_runs(re.sub(r"\s+", " ", keywords).strip())], 9)


def heading1(text_num, text):
    p = para(doc, [], align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4)
    add_runs(p, [{"t": text_num + "."}], 10)
    add_runs(p, [{"t": NBSP + text, "sc": True}], 10)


def heading2(letter, text):
    para(doc, [{"t": "%s.%s%s" % (letter, NBSP, text), "i": True}],
         align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3)


ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII"]
FIGMAP = {"fig2_global": "1", "fig3_dependence": "2",
          "fig4_event_maps": "3", "fig5_cases": "4"}


def add_figure(name, caption, wide=False):
    if wide:
        s = doc.add_section(WD_SECTION.CONTINUOUS)
        page_setup(s)
        set_cols(s, 1, 18)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture("../figures/%s.png" % name, width=FULL_W if wide else COL_W)
    cap = para(doc, [], size=8, before=4, after=10)
    add_runs(cap, [{"t": "Fig.%s%s.%s" % (NBSP, FIGMAP[name], NBSP)}], 8)
    add_runs(cap, tex_runs(re.sub(r"\s+", " ", caption).strip()), 8)
    if wide:
        s = doc.add_section(WD_SECTION.CONTINUOUS)
        page_setup(s)
        set_cols(s, 2, 18)


def add_table_I(caption):
    head = para(doc, [], size=8, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=0)
    add_runs(head, [{"t": "TABLE I"}], 8)
    cap = para(doc, [], size=8, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6)
    add_runs(cap, [dict(r, sc=True) for r in tex_runs(re.sub(r"\s+", " ", caption).strip())], 8)
    rows = [
        ("Regional PM signal (9)", "neighbor PM2.5 mean and count at 25, 50, 100 km; "
         "dispersion at 50 km; CAMS AOD; CAMS PM2.5"),
        ("Wildfire smoke (1)", "HMS smoke tier (none/light/medium/heavy)"),
        ("Meteorology (7)", "temperature, humidity, pressure, wind speed, "
         "precipitation, two interaction terms"),
        ("Geography (4)", "latitude, longitude, distance to coast, distance to "
         "nearest sensor"),
        ("Season and calendar (9)", "month, weekday, day of year, and their "
         "sine/cosine encodings"),
    ]
    t = doc.add_table(rows=len(rows) + 1, cols=2)
    t.style = doc.styles["Normal Table"]
    for c, txt in zip(t.rows[0].cells, ["Group", "Features"]):
        c.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs(c.paragraphs[0], [{"t": txt, "b": True}], 8)
        cell_border(c._tc.get_or_add_tcPr(), "top", "8")
        cell_border(c._tc.get_or_add_tcPr(), "bottom", "4")
    for r, (g, f) in zip(t.rows[1:], rows):
        add_runs(r.cells[0].paragraphs[0], [{"t": g}], 8)
        add_runs(r.cells[1].paragraphs[0], [{"t": f}], 8)
    for c in t.rows[-1].cells:
        cell_border(c._tc.get_or_add_tcPr(), "bottom", "8")
    para(doc, [], size=8, after=6)


content = r"\section" + body.split(r"\section", 1)[1]
content = content.split(r"\bibliographystyle")[0]
content = re.sub(r"(?<!\\)%.*", "", content)
content = content.replace(r"\balance", "")

# Pull floats out of the flow; leave \x01 anchors where they sat.
floats = []
float_pat = r"\\begin\{(table|figure\*?)\}\[[^\]]*\](.*?)\\end\{\1\}"
for m in re.finditer(float_pat, content, re.S):
    kind, inner = m.group(1), m.group(2)
    cap = re.search(r"\\caption\{(.*?)\}\s*\\label", inner, re.S)
    cap = cap.group(1) if cap else ""
    img = re.search(r"\\includegraphics\[[^\]]*\]\{([^}]+)\}", inner)
    floats.append((kind, img.group(1) if img else None, cap))
content = re.sub(float_pat, "\x01", content, flags=re.S)


def emit_float():
    global float_i
    kind, img, cap = floats[float_i]
    float_i += 1
    if img:
        add_figure(img, cap, wide=(kind == "figure*"))
    else:
        add_table_I(cap)


sec_i = sub_i = subsub_i = float_i = 0
tokens = [t for t in re.split(
    r"(\\(?:sub)*section\{[^}]*\}|\\begin\{equation\}.*?\\end\{equation\}|\n\s*\n)",
    content, flags=re.S) if t and t.strip()]

i = 0
while i < len(tokens):
    tok = tokens[i].strip()
    if tok.startswith(r"\section{"):
        sec_i += 1
        sub_i = 0
        heading1(ROMAN[sec_i - 1], re.match(r"\\section\{([^}]*)\}", tok).group(1))
    elif tok.startswith(r"\subsection{"):
        sub_i += 1
        subsub_i = 0
        heading2(chr(64 + sub_i), re.match(r"\\subsection\{([^}]*)\}", tok).group(1))
    elif tok.startswith(r"\subsubsection{"):
        subsub_i += 1
        h = re.match(r"\\subsubsection\{([^}]*)\}", tok).group(1)
        rest = ""
        if i + 1 < len(tokens) and not tokens[i + 1].strip().startswith("\\"):
            rest = tokens[i + 1]
            i += 1
        n_anchor = rest.count("\x01")
        p = para(doc, [], indent=14.4, line=0.95)
        add_runs(p, [{"t": "%d)%s%s:%s" % (subsub_i, NBSP, h, NBSP), "i": True}], 10)
        add_runs(p, tex_runs(rest.replace("\x01", " ").strip()), 10)
        for _ in range(n_anchor):
            emit_float()
    elif tok.startswith(r"\begin{equation}"):
        inner = re.search(r"\\begin\{equation\}(.*?)\\end\{equation\}", tok, re.S).group(1)
        inner = re.sub(r"\\label\{[^}]*\}", "", inner).strip().rstrip(",").strip()
        eq = para(doc, [], align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=12)
        eq.paragraph_format.tab_stops.add_tab_stop(Pt(100), WD_TAB_ALIGNMENT.CENTER)
        eq.paragraph_format.tab_stops.add_tab_stop(Pt(252), WD_TAB_ALIGNMENT.RIGHT)
        eq.add_run("\t").font.size = Pt(10)
        add_runs(eq, math_runs(inner), 10)
        add_runs(eq, [{"t": ","}], 10)
        eq.add_run("\t(1)").font.size = Pt(10)
    else:
        n_anchor = tok.count("\x01")
        text = tok.replace("\x01", " ").strip()
        if text:
            para(doc, tex_runs(text), indent=14.4, line=0.95)
        for _ in range(n_anchor):
            emit_float()
    i += 1

# References (heading centered small caps, entries 8pt hanging-indented).
p = para(doc, [], align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4)
add_runs(p, [{"t": "References", "sc": True}], 10)

items = re.split(r"\\bibitem\{([^}]+)\}", BBL.split(r"\begin{thebibliography}")[1])
for key, txt in zip(items[1::2], items[2::2]):
    txt = txt.split(r"\end{thebibliography}")[0]
    txt = re.sub(r"\\BIBentry[A-Za-z]*", "", txt)
    txt = re.sub(r"\\hskip.*?\\relax", " ", txt)
    txt = txt.replace("\\-", "").replace("--", "\u2013")
    txt = re.sub(r"\\url\{([^}]*)\}", r"\1", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    p = para(doc, [], size=8, after=2.5, line=9, exact=True)
    pf = p.paragraph_format
    pf.left_indent = Pt(14.4)
    pf.first_line_indent = Pt(-14.4)
    add_runs(p, [{"t": "[%s]%s" % (CITE[key], NBSP)}], 8)
    add_runs(p, bbl_runs(txt), 8)

doc.save("main2.docx")
print("wrote main2.docx")
